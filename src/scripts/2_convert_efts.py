"""
2_convert_efts.py
Converte arquivos .pdf e .docx para .txt em "EFTs txt/".
Após cada conversão, extrai também um JSON estruturado para "EFTs json/".
Imagens são transcritas via claude CLI (sem API keys).

Pastas de entrada:
  EFTs Novas PDF/   -> arquivos .pdf  (preferencial)
  EFTs Novas/       -> arquivos .docx (fallback)

Pastas de saída:
  EFTs txt/         -> transcrição completa em texto
  EFTs json/        -> dados técnicos estruturados (para busca por IA)

Uso:
  python src/scripts/2_convert_efts.py [--only "NOME_ARQUIVO"] [--force] [--skip-json]
  --only:      processa apenas o arquivo cujo nome contenha esse trecho
  --force:     reprocessa mesmo se .txt/.json já existirem
  --skip-json: pula a extração de JSON (gera apenas o .txt)
"""
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import unicodedata
import zipfile
from pathlib import Path

ROOT_DIR        = Path(__file__).parent.parent.parent   # nfse-builder/
INPUT_DIR_PDF   = ROOT_DIR / "EFTs Novas PDF"
INPUT_DIR_DOCX  = ROOT_DIR / "EFTs Novas"
OUTPUT_TXT_DIR  = ROOT_DIR / "EFTs txt"
OUTPUT_JSON_DIR = ROOT_DIR / "EFTs json"


# ---------------------------------------------------------------------------
# Claude CLI helpers
# ---------------------------------------------------------------------------

def _clean_env() -> dict:
    """Remove CLAUDECODE para permitir chamada aninhada a partir do Claude Code."""
    env = os.environ.copy()
    env.pop("CLAUDECODE", None)
    return env


def call_claude(prompt: str, timeout: int = 120) -> str:
    """Chama o claude CLI via stdin. Sem API keys."""
    if not shutil.which("claude"):
        raise RuntimeError(
            "claude CLI não encontrado no PATH. "
            "Execute: npm install -g @anthropic-ai/claude-code"
        )
    result = subprocess.run(
        ["claude", "--print"],
        input=prompt, env=_clean_env(),
        capture_output=True, text=True, encoding="utf-8", timeout=timeout
    )
    if result.returncode != 0:
        raise RuntimeError(f"claude CLI erro: {result.stderr[:500]}")
    return result.stdout.strip()


def call_claude_with_image(prompt: str, image_path: str, timeout: int = 60) -> str:
    """Usa claude CLI para transcrever uma imagem via ferramenta Read."""
    if not shutil.which("claude"):
        raise RuntimeError(
            "claude CLI não encontrado no PATH. "
            "Execute: npm install -g @anthropic-ai/claude-code"
        )
    full_prompt = (
        f"{prompt}\n\n"
        f"Use a ferramenta Read para ler o arquivo de imagem: {image_path}"
    )
    result = subprocess.run(
        ["claude", "--print", "--dangerously-skip-permissions", "--tools", "Read"],
        input=full_prompt,
        env=_clean_env(),
        capture_output=True, text=True, encoding="utf-8", timeout=timeout
    )
    if result.returncode != 0:
        raise RuntimeError(f"Erro ao transcrever imagem: {result.stderr[:300]}")
    return result.stdout.strip()


IMAGE_PROMPT = (
    "Esta imagem faz parte de uma especificação funcional (EFT) de NFS-e brasileira. "
    "Use a ferramenta Read para ler o arquivo de imagem indicado abaixo e transcreva TODO o conteúdo visível para texto puro. "
    "Se for tabela, preserve colunas no formato ASCII. "
    "Se for texto ou tela de sistema, transcreva exatamente. "
    "Não adicione comentários — apenas o conteúdo."
)


# ---------------------------------------------------------------------------
# PDF conversion (pymupdf)
# ---------------------------------------------------------------------------

def convert_pdf_to_text(pdf_path: Path) -> str:
    try:
        import fitz
    except ImportError:
        raise RuntimeError(
            "pymupdf não instalado. Execute: pip install pymupdf"
        )

    doc = fitz.open(str(pdf_path))
    lines = []
    image_counter = 0

    for page_num, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            lines.append(f"\n--- Página {page_num + 1} ---")
            lines.append(text)

        image_list = page.get_images(full=True)
        for img_ref in image_list:
            xref = img_ref[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]

            tmp_img = Path(tempfile.mktemp(suffix=f".{image_ext}"))
            tmp_img.write_bytes(image_bytes)

            image_counter += 1
            print(f"    Transcrevendo imagem {image_counter} (pag {page_num + 1})...")
            try:
                transcription = call_claude_with_image(IMAGE_PROMPT, str(tmp_img))
                lines.append(f"\n[IMAGEM {image_counter} TRANSCRITA]")
                lines.append(transcription)
                lines.append(f"[FIM IMAGEM {image_counter}]\n")
            except Exception as e:
                lines.append(f"\n[IMAGEM {image_counter} - erro: {e}]\n")
            finally:
                tmp_img.unlink(missing_ok=True)

    doc.close()
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# DOCX conversion (python-docx + zipfile)
# ---------------------------------------------------------------------------

def _get_docx_images(docx_path: Path) -> dict:
    """Extrai imagens do .docx como {relationship_id: (bytes, ext)}."""
    images = {}
    with zipfile.ZipFile(docx_path, "r") as z:
        import xml.etree.ElementTree as ET
        rels_path = "word/_rels/document.xml.rels"
        rid_to_media = {}
        if rels_path in z.namelist():
            rels_xml = z.read(rels_path)
            tree = ET.fromstring(rels_xml)
            for rel in tree:
                rid = rel.get("Id", "")
                target = rel.get("Target", "")
                if "media/" in target:
                    rid_to_media[rid] = "word/" + target.lstrip("/")

        for rid, media_path in rid_to_media.items():
            if media_path in z.namelist():
                ext = media_path.rsplit(".", 1)[-1].lower()
                images[rid] = (z.read(media_path), ext, media_path.split("/")[-1])

    return images


def convert_docx_to_text(docx_path: Path) -> str:
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        raise RuntimeError(
            "python-docx não instalado. Execute: pip install python-docx"
        )

    doc = Document(docx_path)
    images_data = _get_docx_images(docx_path)

    lines = []
    image_counter = 0

    for para in doc.paragraphs:
        drawings = para._element.findall(".//" + qn("w:drawing"))

        if drawings:
            for drawing in drawings:
                blips = drawing.findall(".//" + qn("a:blip"))
                for blip in blips:
                    embed = blip.get(qn("r:embed"))
                    if embed and embed in images_data:
                        img_bytes, ext, img_name = images_data[embed]
                        tmp_img = Path(tempfile.mktemp(suffix=f".{ext}"))
                        tmp_img.write_bytes(img_bytes)

                        image_counter += 1
                        print(f"    Transcrevendo imagem {image_counter}: {img_name} ...")
                        try:
                            transcription = call_claude_with_image(IMAGE_PROMPT, str(tmp_img))
                            lines.append(f"\n[IMAGEM {image_counter} TRANSCRITA - {img_name}]")
                            lines.append(transcription)
                            lines.append(f"[FIM DA IMAGEM {image_counter}]\n")
                        except Exception as e:
                            lines.append(f"\n[IMAGEM {image_counter} - erro: {e}]\n")
                        finally:
                            tmp_img.unlink(missing_ok=True)
        else:
            text = para.text
            if text.strip():
                lines.append(text)

    for table in doc.tables:
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells))
        if rows_text:
            lines.append("\n" + "\n".join(rows_text) + "\n")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# JSON extraction (estrutura técnica da EFT)
# ---------------------------------------------------------------------------

JSON_SCHEMA = """{
  "tipo_documento": "EFT",
  "cliente": null,
  "projeto": null,
  "municipio": null,
  "estado": null,
  "titulo_programa": null,
  "versao_documento": null,
  "data_documento": null,
  "abrasf_versao": null,
  "transacoes_sap": [],
  "tabelas_sap": [],
  "classes_tecnicas": [],
  "campos_xml": [],
  "regras_negocio": [],
  "regras_tecnicas": [],
  "regras_cancelamento": [],
  "regras_competencia": [],
  "formatos_dados": []
}"""

JSON_EXTRACTION_PROMPT_TEMPLATE = """Analise este documento de EFT (Especificação Funcional-Técnica) de NFS-e e extraia os dados técnicos estruturados.

Regras obrigatórias:
1. Extraia APENAS informações realmente presentes no documento
2. Se um campo não existir, retorne null (para strings) ou [] (para listas)
3. Preserve exatamente a grafia dos nomes técnicos encontrados
4. Não invente dados nem faça suposições
5. Retorne SOMENTE o JSON, sem texto antes ou depois, sem markdown

Campos explicados:
- tipo_documento: sempre "EFT"
- cliente: nome do cliente/empresa contratante (ex: "COGNA")
- projeto: nome do projeto (ex: "NFS-e")
- municipio: nome do município da NFS-e
- estado: sigla do estado (ex: "PR", "SP")
- titulo_programa: título/nome do programa ou EFT
- versao_documento: versão do documento (ex: "1.0", "2.3")
- data_documento: data do documento (ex: "2024-01-15")
- abrasf_versao: versão do padrão ABRASF (ex: "2.04", "2.02", "1.0")
- transacoes_sap: lista de transações SAP mencionadas (ex: ["VF01", "VF02"])
- tabelas_sap: lista de tabelas SAP mencionadas (ex: ["VBRK", "T001"])
- classes_tecnicas: classes ABAP ou técnicas mencionadas
- campos_xml: campos do XML de NFS-e mencionados (ex: ["InfDeclaracaoPrestacaoServico", "Tomador"])
- regras_negocio: lista de regras de negócio relevantes (resumidas, uma por item)
- regras_tecnicas: regras técnicas de implementação (uma por item)
- regras_cancelamento: regras específicas de cancelamento de NFS-e
- regras_competencia: regras de competência/período de apuração
- formatos_dados: formatos especiais de dados mencionados (ex: "CNPJ sem pontuação", "data YYYYMMDD")

Documento EFT:
---
{txt_content}
---

Retorne SOMENTE o JSON (sem markdown, sem explicação):"""


def extract_json_from_txt(txt_content: str, stem: str, out_path: Path, timeout: int = 120) -> bool:
    """
    Chama o claude CLI para extrair o JSON estruturado do conteúdo TXT.
    Salva em out_path. Retorna True se bem-sucedido.
    """
    # Limita tamanho para evitar prompts muito longos (mantém início + fim)
    MAX_CHARS = 40000
    if len(txt_content) > MAX_CHARS:
        half = MAX_CHARS // 2
        txt_trimmed = (
            txt_content[:half]
            + f"\n\n[... {len(txt_content) - MAX_CHARS} caracteres omitidos ...]\n\n"
            + txt_content[-half:]
        )
    else:
        txt_trimmed = txt_content

    prompt = JSON_EXTRACTION_PROMPT_TEMPLATE.format(txt_content=txt_trimmed)

    try:
        raw = call_claude(prompt, timeout=timeout)
    except Exception as e:
        print(f"  [ERRO JSON] claude CLI: {e}")
        return False

    # Extrai JSON do output (remove eventuais blocos markdown)
    raw = re.sub(r"^```(?:json)?\s*\n?", "", raw.strip(), flags=re.MULTILINE)
    raw = re.sub(r"\n?```\s*$", "", raw.strip(), flags=re.MULTILINE)
    raw = raw.strip()

    # Tenta parsear para validar
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"  [ERRO JSON] JSON inválido: {e}")
        # Salva mesmo assim para inspeção
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.with_suffix(".json.invalid").write_text(raw, encoding="utf-8")
        return False

    # Garante que o municipio/estado vêm do nome do arquivo se ausentes
    if not data.get("municipio") or not data.get("estado"):
        # Tenta extrair do stem do arquivo: "COGNA_EFT - NFSe Toledo - PR"
        m = re.search(r"NFSe?\s+(.+?)\s+-\s+([A-Z]{2})", stem, re.IGNORECASE)
        if m:
            if not data.get("municipio"):
                data["municipio"] = m.group(1).strip()
            if not data.get("estado"):
                data["estado"] = m.group(2).upper()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(
        json.dumps(data, ensure_ascii=False, indent=2),
        encoding="utf-8"
    )
    return True


# ---------------------------------------------------------------------------
# Funções de suporte
# ---------------------------------------------------------------------------

def normalize(name: str) -> str:
    nfkd = unicodedata.normalize("NFKD", name)
    return "".join(c for c in nfkd if not unicodedata.combining(c)).lower().strip()


def collect_source_files(only) -> list:
    """
    Retorna lista de (arquivo_fonte, tipo) para processar.
    Prioriza PDF quando o mesmo stem existe nos dois formatos.
    tipo = "pdf" ou "docx"
    """
    pdf_files = {}
    if INPUT_DIR_PDF.exists():
        for f in INPUT_DIR_PDF.glob("*.pdf"):
            pdf_files[normalize(f.stem)] = (f, "pdf")

    docx_files = {}
    if INPUT_DIR_DOCX.exists():
        for f in INPUT_DIR_DOCX.glob("*.docx"):
            docx_files[normalize(f.stem)] = (f, "docx")

    combined = {**docx_files, **pdf_files}
    files = list(combined.values())
    files.sort(key=lambda t: normalize(t[0].stem))

    if only:
        only_norm = normalize(only)
        files = [t for t in files if only_norm in normalize(t[0].stem)]

    return files


def process_file(src_path: Path, file_type: str, force: bool, skip_json: bool) -> bool:
    """
    Converte um arquivo para .txt e opcionalmente extrai .json.
    Retorna True se bem-sucedido.
    """
    out_txt  = OUTPUT_TXT_DIR  / (src_path.stem + ".txt")
    out_json = OUTPUT_JSON_DIR / (src_path.stem + ".json")

    txt_already_existed = out_txt.exists()

    # --- Conversão TXT ---
    if txt_already_existed and not force:
        print(f"  [PULAR TXT] {src_path.name} — .txt já existe.")
        txt_ok = True
    else:
        print(f"  Convertendo ({file_type.upper()}): {src_path.name}")
        try:
            if file_type == "pdf":
                content = convert_pdf_to_text(src_path)
            else:
                content = convert_docx_to_text(src_path)

            OUTPUT_TXT_DIR.mkdir(parents=True, exist_ok=True)
            out_txt.write_text(content, encoding="utf-8")
            print(f"  [OK TXT] {out_txt.name}")
            txt_ok = True
        except Exception as e:
            print(f"  [ERRO TXT] {src_path.name}: {e}")
            txt_ok = False

    # --- Extração JSON ---
    if txt_ok and not skip_json:
        json_already_existed = out_json.exists()
        if json_already_existed and not force and txt_already_existed:
            print(f"  [PULAR JSON] {out_json.name} — já existe.")
        else:
            print(f"  Extraindo JSON: {out_json.name} ...")
            txt_content = out_txt.read_text(encoding="utf-8")
            ok = extract_json_from_txt(txt_content, src_path.stem, out_json)
            if ok:
                print(f"  [OK JSON] {out_json.name}")
            # Falha no JSON não bloqueia o pipeline (txt já foi salvo)

    return txt_ok


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print("=== Etapa 2: Conversão EFT -> TXT + JSON ===\n")
    print(f"  Entrada PDF:  {INPUT_DIR_PDF}")
    print(f"  Entrada DOCX: {INPUT_DIR_DOCX}")
    print(f"  Saida TXT:    {OUTPUT_TXT_DIR}")
    print(f"  Saida JSON:   {OUTPUT_JSON_DIR}\n")

    force     = "--force" in sys.argv
    skip_json = "--skip-json" in sys.argv
    only = None
    if "--only" in sys.argv:
        idx = sys.argv.index("--only")
        if idx + 1 < len(sys.argv):
            only = sys.argv[idx + 1]

    if skip_json:
        print("  [AVISO] --skip-json ativo: JSON nao sera gerado.\n")

    files = collect_source_files(only)

    if not files:
        if only:
            print(f"ERRO: Nenhum arquivo encontrado com '{only}' no nome.")
            sys.exit(1)
        print("Nenhum arquivo .pdf/.docx encontrado nas pastas de entrada.")
        sys.exit(0)

    if only:
        print(f"Filtro ativo: '{only}' -> {len(files)} arquivo(s)\n")
    else:
        print(f"Processando {len(files)} arquivo(s)...\n")

    success = 0
    errors = 0

    for src_path, file_type in files:
        ok = process_file(src_path, file_type, force, skip_json)
        if ok:
            success += 1
        else:
            errors += 1
        print()

    print(f"Concluido: {success} convertido(s), {errors} erro(s).")
    if errors:
        sys.exit(1)


if __name__ == "__main__":
    main()

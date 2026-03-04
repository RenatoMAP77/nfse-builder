"""
2_convert_efts.py
Converte arquivos .docx da pasta EFTs_municipios/ para .txt em EFTs_text/.
Imagens dentro dos .docx são transcritas via GPT-4o Vision da OpenAI.

Uso:
  python 2_convert_efts.py [--only "NOME_ARQUIVO.docx"] [--force]
  --only: processa apenas o arquivo especificado
  --force: reprocessa mesmo se .txt já existir

Variável de ambiente necessária: OPENAI_API_KEY
"""
import base64
import io
import os
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from openai import OpenAI

# Carrega variáveis de ambiente (.env ou prompt interativo)
sys.path.insert(0, str(Path(__file__).parent))
from _env import load_env

BASE_DIR = Path(__file__).parent.parent
INPUT_DIR = BASE_DIR / "EFTs_municipios"
OUTPUT_DIR = BASE_DIR / "EFTs_text"


def get_images_from_docx(docx_path: Path) -> dict[str, bytes]:
    """
    Extrai todas as imagens do .docx como dicionário {relationship_id: bytes}.
    O arquivo .docx é um ZIP com imagens em word/media/.
    """
    images = {}
    with zipfile.ZipFile(docx_path, "r") as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                images[name] = z.read(name)
    return images


def transcribe_image_with_openai(client: OpenAI, image_bytes: bytes, img_name: str) -> str:
    """Usa GPT-4o Vision para transcrever o conteúdo de uma imagem."""
    # Detecta o tipo da imagem pelo nome
    ext = img_name.rsplit(".", 1)[-1].lower() if "." in img_name else "png"
    mime_map = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png",
                "gif": "image/gif", "bmp": "image/bmp", "webp": "image/webp"}
    mime_type = mime_map.get(ext, "image/png")

    b64 = base64.b64encode(image_bytes).decode("utf-8")
    data_uri = f"data:{mime_type};base64,{b64}"

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Esta imagem faz parte de uma especificação funcional técnica (EFT) "
                            "de integração NFS-e (Nota Fiscal de Serviços Eletrônica) com prefeitura brasileira. "
                            "Transcreva TODO o conteúdo visível desta imagem para texto puro, "
                            "mantendo tabelas no formato ASCII com colunas alinhadas. "
                            "Se for uma tabela, preserve todos os campos, valores e descrições. "
                            "Se for texto, transcreva exatamente. "
                            "Não adicione comentários ou explicações — apenas o conteúdo transcrito."
                        ),
                    },
                    {
                        "type": "image_url",
                        "image_url": {"url": data_uri, "detail": "high"},
                    },
                ],
            }
        ],
        max_tokens=2000,
    )
    return response.choices[0].message.content.strip()


def build_image_remap(docx_path: Path) -> dict[str, str]:
    """
    Constrói mapeamento de relationship id → nome do arquivo de imagem.
    Ex: { "rId5": "word/media/image1.png" }
    """
    remap = {}
    with zipfile.ZipFile(docx_path, "r") as z:
        rels_path = "word/_rels/document.xml.rels"
        if rels_path in z.namelist():
            import xml.etree.ElementTree as ET
            rels_xml = z.read(rels_path)
            tree = ET.fromstring(rels_xml)
            for rel in tree:
                rid = rel.get("Id", "")
                target = rel.get("Target", "")
                if "media/" in target:
                    remap[rid] = "word/" + target.lstrip("/")
    return remap


def convert_docx_to_text(docx_path: Path, client: OpenAI) -> str:
    """
    Converte um .docx para texto, transcrevendo imagens com GPT-4o Vision.
    Retorna o conteúdo como string.
    """
    doc = Document(docx_path)
    images_data = get_images_from_docx(docx_path)
    image_remap = build_image_remap(docx_path)

    lines = []
    image_counter = 0

    for para in doc.paragraphs:
        # Verifica se o parágrafo contém imagens (drawing elements)
        drawings = para._element.findall(
            ".//" + qn("w:drawing")
        )

        if drawings:
            for drawing in drawings:
                # Busca o blip (referência da imagem)
                blips = drawing.findall(
                    ".//" + qn("a:blip")
                )
                for blip in blips:
                    embed = blip.get(qn("r:embed"))
                    if embed and embed in image_remap:
                        img_path = image_remap[embed]
                        if img_path in images_data:
                            image_counter += 1
                            img_bytes = images_data[img_path]
                            img_name = img_path.split("/")[-1]
                            print(f"    Transcrevendo imagem {image_counter}: {img_name} ...")
                            try:
                                transcription = transcribe_image_with_openai(
                                    client, img_bytes, img_name
                                )
                                lines.append(f"\n[IMAGEM {image_counter} TRANSCRITA - {img_name}]")
                                lines.append(transcription)
                                lines.append(f"[FIM DA IMAGEM {image_counter}]\n")
                            except Exception as e:
                                lines.append(f"\n[IMAGEM {image_counter} - erro na transcrição: {e}]\n")
        else:
            text = para.text
            if text.strip():
                lines.append(text)

    # Também processa tabelas do documento
    for table in doc.tables:
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells))
        if rows_text:
            lines.append("\n" + "\n".join(rows_text) + "\n")

    return "\n".join(lines)


def get_output_filename(docx_path: Path) -> Path:
    """Retorna o caminho do .txt correspondente ao .docx."""
    return OUTPUT_DIR / (docx_path.stem + ".txt")


def process_file(docx_path: Path, client: OpenAI, force: bool) -> bool:
    """Processa um único arquivo .docx. Retorna True se sucesso."""
    out_path = get_output_filename(docx_path)

    if out_path.exists() and not force:
        print(f"  [PULAR] {docx_path.name} — .txt já existe. Use --force para reprocessar.")
        return True

    print(f"  Convertendo: {docx_path.name}")
    try:
        content = convert_docx_to_text(docx_path, client)
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        out_path.write_text(content, encoding="utf-8")
        print(f"  [OK] Salvo em: {out_path.name}")
        return True
    except Exception as e:
        print(f"  [ERRO] {docx_path.name}: {e}")
        return False


def main():
    print("=== Etapa 2: Conversão DOCX → TXT ===\n")
    load_env(["OPENAI_API_KEY"])

    client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])

    force = "--force" in sys.argv
    only = None
    if "--only" in sys.argv:
        idx = sys.argv.index("--only")
        if idx + 1 < len(sys.argv):
            only = sys.argv[idx + 1]

    if not INPUT_DIR.exists():
        print(f"ERRO: Pasta {INPUT_DIR} não encontrada.")
        sys.exit(1)

    docx_files = list(INPUT_DIR.glob("*.docx"))
    if not docx_files:
        print(f"Nenhum arquivo .docx encontrado em {INPUT_DIR}")
        sys.exit(0)

    if only:
        docx_files = [f for f in docx_files if only in f.name]
        if not docx_files:
            print(f"ERRO: Nenhum arquivo encontrado com '{only}' no nome.")
            sys.exit(1)

    print(f"Processando {len(docx_files)} arquivo(s)...\n")
    success = 0
    errors = 0

    for docx_path in sorted(docx_files):
        ok = process_file(docx_path, client, force)
        if ok:
            success += 1
        else:
            errors += 1

    print(f"\nConcluído: {success} convertidos, {errors} erros.")
    if errors:
        sys.exit(1)


if __name__ == "__main__":
    main()
